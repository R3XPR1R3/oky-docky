"""
Улучшенный модуль для обнаружения и обработки плейсхолдеров в документах
Поддерживает PDF, DOCX, TXT файлы с надежным извлечением плейсхолдеров
"""

import re
import logging
from typing import List, Dict, Set, Optional
from pathlib import Path
import PyPDF2
import docx
from dataclasses import dataclass

# Настройка логирования
logger = logging.getLogger(__name__)

@dataclass
class PlaceholderInfo:
    """Информация о найденном плейсхолдере"""
    name: str
    count: int
    positions: List[int]
    file_type: str

class AdvancedPlaceholderParser:
    """
    Продвинутый парсер плейсхолдеров с поддержкой разных форматов файлов
    """
    
    def __init__(self):
        # Паттерн для поиска плейсхолдеров: {name}, {{name}}, [name], <name>
        self.patterns = [
            r'\{([a-zA-Z_][a-zA-Z0-9_]*)\}',          # {name}
            r'\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}',      # {{name}}
            r'\[([a-zA-Z_][a-zA-Z0-9_]*)\]',          # [name]
            r'<([a-zA-Z_][a-zA-Z0-9_]*)>',            # <name>
            r'\$\{([a-zA-Z_][a-zA-Z0-9_]*)\}',        # ${name}
        ]
        
        # Скомпилированные регулярные выражения
        self.compiled_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in self.patterns]
        
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Безопасное извлечение текста из PDF файла
        """
        try:
            text_content = ""
            
            # Пробуем PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Ошибка извлечения текста со страницы {page_num}: {e}")
                        continue
            
            # Если PyPDF2 не дал результатов, можно попробовать другие библиотеки
            if not text_content.strip():
                logger.warning(f"PyPDF2 не смог извлечь текст из {file_path}")
                # Здесь можно добавить pdfplumber или другие библиотеки
            
            return text_content
            
        except Exception as e:
            logger.error(f"Критическая ошибка при чтении PDF {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Извлечение текста из DOCX файла
        """
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
                
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                text_content += "\n"
                        
            return text_content
            
        except Exception as e:
            logger.error(f"Ошибка при чтении DOCX {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """
        Извлечение текста из обычного текстового файла
        """
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
                    
            logger.warning(f"Не удалось определить кодировку файла {file_path}")
            return ""
            
        except Exception as e:
            logger.error(f"Ошибка при чтении TXT {file_path}: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: Path) -> str:
        """
        Универсальный метод извлечения текста из файла
        """
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension in ['.txt', '.text']:
            return self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат файла: {file_extension}")
            return ""
    
    def find_placeholders_in_text(self, text: str) -> Dict[str, PlaceholderInfo]:
        """
        Поиск всех плейсхолдеров в тексте
        """
        placeholders = {}
        
        for pattern in self.compiled_patterns:
            matches = list(pattern.finditer(text))
            
            for match in matches:
                placeholder_name = match.group(1).lower()  # Приводим к нижнему регистру для унификации
                
                if placeholder_name not in placeholders:
                    placeholders[placeholder_name] = PlaceholderInfo(
                        name=placeholder_name,
                        count=0,
                        positions=[],
                        file_type="text"
                    )
                
                placeholders[placeholder_name].count += 1
                placeholders[placeholder_name].positions.append(match.start())
        
        return placeholders
    
    def analyze_template_file(self, file_path: Path) -> Dict[str, PlaceholderInfo]:
        """
        Полный анализ файла шаблона с извлечением всех плейсхолдеров
        """
        logger.info(f"🔍 Анализ шаблона: {file_path.name}")
        
        # Извлекаем текст
        text_content = self.extract_text_from_file(file_path)
        
        if not text_content:
            logger.warning(f"❌ Не удалось извлечь текст из {file_path.name}")
            return {}
        
        logger.debug(f"📝 Извлечено {len(text_content)} символов текста")
        logger.debug(f"Начало текста: {text_content[:200]}...")
        
        # Находим плейсхолдеры
        placeholders = self.find_placeholders_in_text(text_content)
        
        # Обновляем тип файла
        file_extension = file_path.suffix.lower()
        for placeholder_info in placeholders.values():
            placeholder_info.file_type = file_extension
        
        if placeholders:
            logger.info(f"✅ Найдено плейсхолдеров: {len(placeholders)}")
            for name, info in placeholders.items():
                logger.info(f"   - {{{name}}} (встречается {info.count} раз)")
        else:
            logger.warning(f"⚠️ Плейсхолдеры не найдены в {file_path.name}")
            
        return placeholders
    
    def scan_templates_directory(self, templates_dir: Path) -> Dict[str, Dict[str, PlaceholderInfo]]:
        """
        Сканирование директории с шаблонами
        """
        logger.info(f"📂 Сканирование директории шаблонов: {templates_dir}")
        
        results = {}
        supported_extensions = ['.pdf', '.docx', '.txt']
        
        for file_path in templates_dir.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    placeholders = self.analyze_template_file(file_path)
                    if placeholders:
                        results[str(file_path)] = placeholders
                except Exception as e:
                    logger.error(f"❌ Ошибка анализа файла {file_path.name}: {e}")
        
        logger.info(f"🎉 Обработано файлов: {len(results)}")
        return results
    
    def generate_form_fields(self, placeholders: Dict[str, PlaceholderInfo]) -> List[Dict]:
        """
        Генерация полей формы на основе найденных плейсхолдеров
        """
        form_fields = []
        
        for name, info in placeholders.items():
            # Определяем тип поля на основе имени плейсхолдера
            field_type = self._detect_field_type(name)
            
            field = {
                'name': name,
                'label': self._generate_label(name),
                'type': field_type,
                'required': True,
                'placeholder': f'Введите {self._generate_label(name).lower()}',
                'validation': self._get_validation_rules(field_type),
                'occurrences': info.count
            }
            
            form_fields.append(field)
        
        # Сортируем поля по приоритету
        form_fields.sort(key=lambda x: self._get_field_priority(x['name']))
        
        return form_fields
    
    def _detect_field_type(self, field_name: str) -> str:
        """Определение типа поля на основе его имени"""
        field_name = field_name.lower()
        
        if any(keyword in field_name for keyword in ['email', 'mail', 'почта']):
            return 'email'
        elif any(keyword in field_name for keyword in ['phone', 'tel', 'телефон']):
            return 'tel'
        elif any(keyword in field_name for keyword in ['date', 'дата']):
            return 'date'
        elif any(keyword in field_name for keyword in ['number', 'num', 'номер', 'количество']):
            return 'number'
        elif any(keyword in field_name for keyword in ['url', 'website', 'сайт']):
            return 'url'
        else:
            return 'text'
    
    def _generate_label(self, field_name: str) -> str:
        """Генерация человекочитаемой метки для поля"""
        # Словарь для перевода распространенных названий
        translations = {
            'name': 'Имя',
            'firstname': 'Имя',
            'lastname': 'Фамилия',
            'email': 'Email',
            'phone': 'Телефон',
            'date': 'Дата',
            'address': 'Адрес',
            'city': 'Город',
            'company': 'Компания',
            'position': 'Должность',
            'test': 'Тест',
            'data': 'Данные'
        }
        
        if field_name.lower() in translations:
            return translations[field_name.lower()]
        
        # Если нет перевода, форматируем название
        return field_name.replace('_', ' ').title()
    
    def _get_validation_rules(self, field_type: str) -> Dict:
        """Получение правил валидации для типа поля"""
        rules = {
            'email': {'pattern': r'^[^\s@]+@[^\s@]+\.[^\s@]+$', 'message': 'Введите корректный email'},
            'tel': {'pattern': r'^[\+]?[\d\s\-\(\)]+$', 'message': 'Введите корректный номер телефона'},
            'url': {'pattern': r'^https?:\/\/.+', 'message': 'Введите корректный URL'},
            'number': {'pattern': r'^\d+$', 'message': 'Введите число'}
        }
        
        return rules.get(field_type, {})
    
    def _get_field_priority(self, field_name: str) -> int:
        """Определение приоритета поля для сортировки"""
        priorities = {
            'name': 1, 'firstname': 1, 'lastname': 2,
            'email': 3, 'phone': 4, 'date': 5,
            'company': 6, 'position': 7
        }
        
        return priorities.get(field_name.lower(), 100)


# Пример использования
def main():
    """
    Пример использования парсера плейсхолдеров
    """
    # Создаем парсер
    parser = AdvancedPlaceholderParser()
    
    # Анализируем конкретный файл
    template_path = Path("templates/document_template.pdf")
    if template_path.exists():
        placeholders = parser.analyze_template_file(template_path)
        
        # Генерируем поля формы
        form_fields = parser.generate_form_fields(placeholders)
        
        print("📋 Сгенерированные поля формы:")
        for field in form_fields:
            print(f"  - {field['label']} ({field['type']})")
    
    # Сканируем всю директорию
    templates_dir = Path("templates")
    if templates_dir.exists():
        all_results = parser.scan_templates_directory(templates_dir)
        
        print(f"\n📊 Общая статистика:")
        print(f"  - Обработано файлов: {len(all_results)}")
        
        total_placeholders = sum(len(placeholders) for placeholders in all_results.values())
        print(f"  - Найдено уникальных плейсхолдеров: {total_placeholders}")


if __name__ == "__main__":
    # Настройка логирования для тестирования
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    main()